from __future__ import annotations

import re
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "research" / "Lyne_Next_Step_Decision_and_Funding_Dossier_2026-07-31.md"
OUTPUT = ROOT / "outputs" / "Lyne_Next_Step_Decision_and_Funding_Dossier_2026-07-31.docx"

# decision_memo preset: Arial, compact professional hierarchy, 6 pt body spacing,
# H1 12/6 and H2 10/5 before/after, restrained color.
NAVY = "17324D"
TEAL = "087E8B"
PALE_TEAL = "E8F4F5"
PALE_BLUE = "EDF3F8"
LIGHT_GREY = "F2F4F6"
MID_GREY = "6A7580"
DARK = "17212B"
WHITE = "FFFFFF"
RED = "A33A3A"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=95, bottom=80, end=95) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color="D7DEE4", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_hyperlink(paragraph, text: str, url: str, color=TEAL, underline=True):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    run_props.append(color_node)
    if underline:
        underline_node = OxmlElement("w:u")
        underline_node.set(qn("w:val"), "single")
        run_props.append(underline_node)
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|https?://[^\s)]+)"
)


def add_inline(paragraph, text: str, default_bold=False, default_color=None) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            run.bold = default_bold
            if default_color:
                run.font.color.rgb = RGBColor.from_string(default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            if default_color:
                run.font.color.rgb = RGBColor.from_string(default_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(NAVY)
            run.font.highlight_color = None
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            url = token.rstrip(".,;")
            trailing = token[len(url) :]
            display = urlparse(url).netloc or url
            add_hyperlink(paragraph, display, url)
            if trailing:
                paragraph.add_run(trailing)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        run.bold = default_bold
        if default_color:
            run.font.color.rgb = RGBColor.from_string(default_color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_bottom_rule(paragraph, color=TEAL, size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.widow_control = True

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    styles["Title"].font.size = Pt(29)
    styles["Title"].font.bold = True
    styles["Title"].paragraph_format.space_after = Pt(12)

    styles["Subtitle"].font.size = Pt(12)
    styles["Subtitle"].font.color.rgb = RGBColor.from_string(MID_GREY)
    styles["Subtitle"].paragraph_format.space_after = Pt(8)

    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)

    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(TEAL)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)

    styles["Heading 3"].font.size = Pt(10.5)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    styles["Heading 4"].font.size = Pt(9.5)
    styles["Heading 4"].font.bold = True
    styles["Heading 4"].paragraph_format.space_before = Pt(6)
    styles["Heading 4"].paragraph_format.space_after = Pt(3)

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Arial"
        style.font.size = Pt(9.2)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)


def configure_section(section) -> None:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)


def add_header_footer(section, first=False) -> None:
    section.different_first_page_header_footer = first
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Lyne  •  Decision & Funding Dossier")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MID_GREY)
    add_bottom_rule(p, color="D7DEE4", size="4")

    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    left = p.add_run("Prepared 31 July 2026  •  Planning document, not a bid")
    left.font.name = "Arial"
    left.font.size = Pt(7.5)
    left.font.color.rgb = RGBColor.from_string(MID_GREY)
    p.add_run("\t")
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    add_page_number(p)
    if first:
        section.first_page_header.paragraphs[0].clear()
        section.first_page_footer.paragraphs[0].clear()


def add_cover(doc: Document, bottom_line: str) -> None:
    section = doc.sections[0]
    configure_section(section)
    section.top_margin = Cm(1.45)
    add_header_footer(section, first=True)

    mast = doc.add_table(rows=1, cols=2)
    mast.alignment = WD_TABLE_ALIGNMENT.CENTER
    mast.autofit = False
    mast.columns[0].width = Cm(8)
    mast.columns[1].width = Cm(9.5)
    left, right = mast.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, 40, 20, 40, 20)
    p = left.paragraphs[0]
    run = p.add_run("Lyne")
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("DECISION & FUNDING DOSSIER")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    title = doc.add_paragraph(style="Title")
    title.add_run("Lyne:\nNext-Step Decision\nand Funding Dossier")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Market choice, court positioning, product readiness, funding routes and launch budgets")

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta.columns[0].width = Cm(3.2)
    meta.columns[1].width = Cm(13.8)
    values = [
        ("DECISION DATE", "31 July 2026"),
        ("PREPARED FOR", "Lyne founder and family funding discussion"),
        ("GEOGRAPHY", "Jamaica, with regional-scale architecture"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        set_cell_shading(row.cells[0], PALE_BLUE)
        for cell in row.cells:
            set_cell_margins(cell, 80, 100, 80, 100)
        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        r.font.name = "Arial"
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p = row.cells[1].paragraphs[0]
        add_inline(p, value)
    set_table_borders(meta, color="E0E5E9", size="3")

    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    callout.autofit = False
    callout.columns[0].width = Cm(17)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, PALE_TEAL)
    set_cell_margins(cell, 170, 190, 170, 190)
    p = cell.paragraphs[0]
    r = p.add_run("RECOMMENDED DECISION\n")
    r.font.name = "Arial"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    add_inline(p, bottom_line)
    set_table_borders(callout, color="B7DDE0", size="5")

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(16)
    note.paragraph_format.space_after = Pt(0)
    r = note.add_run(
        "Evidence cut-off: 31 July 2026. Prices and open calls must be rechecked before payment or submission."
    )
    r.font.name = "Arial"
    r.font.size = Pt(7.8)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(MID_GREY)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph("How to use this dossier", style="Heading 1")
    add_bottom_rule(p, color=TEAL, size="7")
    intro = doc.add_paragraph()
    add_inline(
        intro,
        "Read Sections 1 and 4 for the decision, Section 10 for funding, Section 12 for the budget, and Section 13 for the execution plan. The remaining sections preserve the evidence and decision logic needed for prospect and grant conversations.",
    )
    items = [
        "1–4  Decision, product readiness and market ranking",
        "5–7  Financial services, courts, universities and diagnostics",
        "8–9  Personas, demo design and targeting",
        "10–11  Grants, incorporation, data protection and procurement",
        "12  Full private and public-enterprise budgets",
        "13–14  90-day roadmap and go/no-go gates",
        "15  Source register",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, item)
    doc.add_page_break()


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    set_table_fixed_layout(table)
    set_repeat_table_header(table.rows[0])

    for row_index, (word_row, doc_row) in enumerate(zip(rows, table.rows)):
        prevent_row_split(doc_row)
        if row_index == 0:
            doc_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for col_index, (value, cell) in enumerate(zip(word_row, doc_row.cells)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GREY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_together = True
            add_inline(
                p,
                value,
                default_bold=(row_index == 0),
                default_color=(WHITE if row_index == 0 else None),
            )
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(7.4 if width >= 5 else 8.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_TEAL)
    set_cell_margins(cell, 130, 150, 130, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text)
    set_table_borders(table, color="B7DDE0", size="5")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    add_bottom_rule(p, color="D7DEE4", size="4")


def render_markdown(doc: Document, source_text: str) -> None:
    lines = source_text.splitlines()
    # Cover metadata occupies the beginning; body starts at Section 1.
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1. "))
    lines = lines[start:]

    paragraph_buffer: list[str] = []
    index = 0

    def flush_paragraph():
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        p = doc.add_paragraph()
        add_inline(p, text)
        paragraph_buffer = []

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped == "---":
            flush_paragraph()
            add_horizontal_rule(doc)
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            raw_table: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                raw_table.append(lines[index].strip())
                index += 1
            parsed = [split_table_row(row) for row in raw_table if not is_table_separator(row)]
            add_table(doc, parsed)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1)) - 1
            text = heading_match.group(2)
            style = f"Heading {level}"
            p = doc.add_paragraph(style=style)
            add_inline(p, text)
            if level == 1:
                add_bottom_rule(p, color=TEAL, size="6")
            index += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            quote_parts = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_parts.append(lines[index].strip().lstrip(">").strip())
                index += 1
            add_callout(doc, " ".join(quote_parts))
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.+)$", line)
        if bullet_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet_match.group(2))
            index += 1
            continue

        number_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if number_match:
            flush_paragraph()
            source_number = re.match(r"^\s*(\d+)\.", line).group(1)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(3)
            number_run = p.add_run(f"{source_number}. ")
            number_run.bold = True
            number_run.font.color.rgb = RGBColor.from_string(NAVY)
            add_inline(p, number_match.group(2))
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def add_document_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Lyne: Next-Step Decision and Funding Dossier"
    props.subject = "Jamaica market strategy, funding, procurement and launch budget"
    props.author = "Lyne"
    props.keywords = "Lyne, Jamaica, queue management, credit unions, courts, grants, budget"
    props.comments = "Evidence cut-off 31 July 2026"


def build() -> None:
    source_text = SOURCE.read_text(encoding="utf-8")
    bottom_line_match = re.search(r'^> \*\*Bottom line:\*\* (.+)$', source_text, re.MULTILINE)
    bottom_line = (
        bottom_line_match.group(1)
        if bottom_line_match
        else "Lead with a private credit-union pilot and keep court work as an integration-led discovery track."
    )

    doc = Document()
    configure_styles(doc)
    add_document_properties(doc)
    add_cover(doc, bottom_line)
    add_toc(doc)
    render_markdown(doc, source_text)

    for section_index, section in enumerate(doc.sections):
        configure_section(section)
        add_header_footer(section, first=(section_index == 0))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
