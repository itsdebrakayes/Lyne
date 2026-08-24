from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_lyne_dossier import (
    DARK,
    LIGHT_GREY,
    MID_GREY,
    NAVY,
    PALE_BLUE,
    PALE_TEAL,
    TEAL,
    WHITE,
    add_bottom_rule,
    add_inline,
    add_page_number,
    is_table_separator,
    prevent_row_split,
    set_cell_margins,
    set_cell_shading,
    set_repeat_table_header,
    set_table_borders,
    set_table_fixed_layout,
    split_table_row,
)


ROOT = Path(__file__).resolve().parents[1]
RESEARCH = ROOT / "research"
OUTPUTS = ROOT / "outputs"

DOCS = [
    {
        "source": RESEARCH / "Lyne_App_Only_Budget_2026-07-31.md",
        "output": OUTPUTS / "Lyne_App_Only_Budget_2026-07-31.docx",
        "title": "Lyne App-Only\nLaunch Budget",
        "subtitle": "What the software costs, why every line is necessary, and what can wait",
        "short_title": "App-Only Launch Budget",
        "document_type": "FAMILY FUNDING BRIEF",
        "prepared_for": "Founder and family funding discussion",
        "decision": (
            "Plan around J$5.66M (US$35.38k) for a responsible app-only launch, "
            "or J$2.44M (US$15.26k) if the founder completes the remaining product work."
        ),
        "subject": "Detailed app-only launch budget in USD and JMD",
        "one_page": False,
    },
    {
        "source": RESEARCH / "Lyne_Granular_Feasibility_Study_2026-07-31.md",
        "output": OUTPUTS / "Lyne_Granular_Feasibility_Study_2026-07-31.docx",
        "title": "Lyne Granular\nFeasibility Study",
        "subtitle": "Cost, price, customer capacity, break-even, delivery risk and go/no-go gates",
        "short_title": "Granular Feasibility Study",
        "document_type": "INVESTMENT DECISION",
        "prepared_for": "Founder, family and prospective funders",
        "decision": (
            "Conditionally feasible: win and measure a paid private pilot before completing "
            "the full commercial build or depending on a court contract."
        ),
        "subject": "Detailed commercial and operational feasibility analysis",
        "one_page": False,
    },
    {
        "source": RESEARCH / "Lyne_Feasibility_One_Page_Decision_2026-07-31.md",
        "output": OUTPUTS / "Lyne_Feasibility_One_Page_Decision_2026-07-31.docx",
        "title": "Lyne: One-Page Investment Decision",
        "subtitle": "The cost, commercial test and clear answer",
        "short_title": "One-Page Investment Decision",
        "document_type": "DECISION IN ONE PAGE",
        "prepared_for": "Founder and family",
        "decision": (
            "YES—conditionally. Proceed in stages, win a paid private pilot, and do not spend "
            "J$11M before customer proof."
        ),
        "subject": "One-page Lyne feasibility decision",
        "one_page": True,
    },
]


def configure_section(section, one_page: bool = False) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    if one_page:
        # The one-page brief keeps the same decision-memo page geometry.
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.72)


def configure_styles(doc: Document, one_page: bool = False) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.25 if one_page else 9.25)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(2.2 if one_page else 6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.widow_control = True

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    styles["Title"].font.size = Pt(22 if one_page else 29)
    styles["Title"].font.bold = True
    styles["Title"].paragraph_format.space_after = Pt(4 if one_page else 10)

    styles["Subtitle"].font.size = Pt(9 if one_page else 11.5)
    styles["Subtitle"].font.color.rgb = RGBColor.from_string(MID_GREY)
    styles["Subtitle"].paragraph_format.space_after = Pt(4 if one_page else 8)

    styles["Heading 1"].font.size = Pt(10.5 if one_page else 16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].paragraph_format.space_before = Pt(4 if one_page else 12)
    styles["Heading 1"].paragraph_format.space_after = Pt(2 if one_page else 6)

    styles["Heading 2"].font.size = Pt(9.3 if one_page else 12.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(TEAL)
    styles["Heading 2"].paragraph_format.space_before = Pt(3 if one_page else 9)
    styles["Heading 2"].paragraph_format.space_after = Pt(2 if one_page else 4)

    styles["Heading 3"].font.size = Pt(8.7 if one_page else 10.5)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].paragraph_format.space_before = Pt(2 if one_page else 7)
    styles["Heading 3"].paragraph_format.space_after = Pt(1.5 if one_page else 3)

    styles["Heading 4"].font.size = Pt(8.4 if one_page else 9.5)
    styles["Heading 4"].font.bold = True
    styles["Heading 4"].paragraph_format.space_before = Pt(2 if one_page else 5)
    styles["Heading 4"].paragraph_format.space_after = Pt(1 if one_page else 2)

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Arial"
        style.font.size = Pt(7.9 if one_page else 9.1)
        style.paragraph_format.left_indent = Inches(0.22)
        style.paragraph_format.first_line_indent = Inches(-0.16)
        style.paragraph_format.space_after = Pt(1.2 if one_page else 3)


def set_table_width(table, total_width_twips: int = 9240, indent_twips: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_twips))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def new_numbering_instance(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    num_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    # Built-in "List Number" uses abstract numbering definition 7.
    abstract.set(qn("w:val"), "7")
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def column_widths(rows: list[list[str]], total: int = 9240) -> list[int]:
    count = max(len(row) for row in rows)
    first_header = rows[0][0].strip() if rows and rows[0] else ""
    if first_header == "#" and count == 5:
        return [420, 1900, 1350, 1650, 3920]
    if count == 1:
        return [total]
    if count == 2:
        ratios = [0.38, 0.62]
    elif count == 3:
        ratios = [0.48, 0.23, 0.29]
    elif count == 4:
        ratios = [0.42, 0.19, 0.19, 0.20]
    elif count == 5:
        ratios = [0.34, 0.165, 0.165, 0.165, 0.165]
    else:
        ratios = [1 / count] * count
    widths = [int(total * ratio) for ratio in ratios]
    widths[-1] += total - sum(widths)
    return widths


def add_header_footer(section, short_title: str, first_page_blank: bool) -> None:
    section.different_first_page_header_footer = first_page_blank

    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Lyne  •  {short_title}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MID_GREY)
    add_bottom_rule(p, color="D7DEE4", size="4")

    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    left = p.add_run("Prepared 31 July 2026  •  Planning document")
    left.font.name = "Arial"
    left.font.size = Pt(7.3)
    left.font.color.rgb = RGBColor.from_string(MID_GREY)
    p.add_run("\t")
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.35))
    add_page_number(p)

    if first_page_blank:
        section.first_page_header.paragraphs[0].clear()
        section.first_page_footer.paragraphs[0].clear()


def add_masthead(doc: Document, document_type: str, compact: bool = False) -> None:
    mast = doc.add_table(rows=1, cols=2)
    mast.alignment = WD_TABLE_ALIGNMENT.CENTER
    mast.autofit = False
    set_table_fixed_layout(mast)
    set_table_width(mast)
    widths = [4200, 5040]
    for index, width in enumerate(widths):
        mast.columns[index].width = Inches(width / 1440)
        set_cell_width(mast.rows[0].cells[index], width)

    left, right = mast.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, 25 if compact else 40, 20, 25 if compact else 40, 20)
    p = left.paragraphs[0]
    run = p.add_run("Lyne")
    run.font.name = "Arial"
    run.font.size = Pt(15 if compact else 18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(document_type)
    run.font.name = "Arial"
    run.font.size = Pt(8 if compact else 9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    add_bottom_rule(right.paragraphs[0], color="D7DEE4", size="4")


def add_callout(doc: Document, text: str, compact: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed_layout(table)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_width(cell, 9240)
    set_cell_shading(cell, PALE_TEAL)
    set_cell_margins(cell, 80 if compact else 130, 130, 80 if compact else 130, 130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text)
    for run in p.runs:
        run.font.size = Pt(8 if compact else 9.4)
    set_table_borders(table, color="B7DDE0", size="5")


def add_cover(doc: Document, spec: dict) -> None:
    add_masthead(doc, spec["document_type"])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(44)

    title = doc.add_paragraph(style="Title")
    title.add_run(spec["title"])
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(spec["subtitle"])

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_table_fixed_layout(meta)
    set_table_width(meta)
    widths = [2100, 7140]
    values = [
        ("PREPARED", "31 July 2026"),
        ("PREPARED FOR", spec["prepared_for"]),
        ("PLANNING RATE", "US$1 = J$160"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        set_cell_width(row.cells[0], widths[0])
        set_cell_width(row.cells[1], widths[1])
        set_cell_shading(row.cells[0], PALE_BLUE)
        for cell in row.cells:
            set_cell_margins(cell, 70, 90, 70, 90)
        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        r.font.name = "Arial"
        r.font.size = Pt(7.2)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p = row.cells[1].paragraphs[0]
        add_inline(p, value)
    set_table_borders(meta, color="E0E5E9", size="3")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)
    add_callout(doc, spec["decision"])

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(14)
    r = note.add_run(
        "Evidence cut-off: 31 July 2026. Vendor prices and eligibility should be reconfirmed before payment."
    )
    r.font.name = "Arial"
    r.font.size = Pt(7.7)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(MID_GREY)
    doc.add_page_break()


def add_one_page_opening(doc: Document, spec: dict) -> None:
    add_masthead(doc, spec["document_type"], compact=True)
    title = doc.add_paragraph(style="Title")
    title.add_run(spec["title"].replace("\n", " "))
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(f"{spec['subtitle']}  •  31 July 2026  •  US$1 = J$160")
    add_callout(doc, spec["decision"], compact=True)


def add_table(doc: Document, rows: list[list[str]], one_page: bool = False) -> None:
    if not rows:
        return
    count = max(len(row) for row in rows)
    rows = [row + [""] * (count - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed_layout(table)
    set_table_width(table)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    widths = column_widths(rows)

    for row_index, (source_row, doc_row) in enumerate(zip(rows, table.rows)):
        prevent_row_split(doc_row)
        if row_index == 0:
            doc_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for col_index, (value, cell) in enumerate(zip(source_row, doc_row.cells)):
            set_cell_width(cell, widths[col_index])
            cell.width = Inches(widths[col_index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(
                cell,
                top=38 if one_page else 68,
                start=55 if one_page else 82,
                bottom=38 if one_page else 68,
                end=55 if one_page else 82,
            )
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GREY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_together = True
            add_inline(
                p,
                value,
                default_bold=(row_index == 0),
                default_color=(WHITE if row_index == 0 else None),
            )
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(
                    6.7 if one_page else (7.2 if count >= 5 else 7.8)
                )
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0 if one_page else 1)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    add_bottom_rule(p, color="D7DEE4", size="4")


def render_markdown(doc: Document, source_text: str, one_page: bool = False) -> None:
    lines = source_text.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## "))
    lines = lines[start:]
    paragraph_buffer: list[str] = []
    index = 0
    current_numbering_id: int | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        p = doc.add_paragraph()
        add_inline(p, text)
        paragraph_buffer = []

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            current_numbering_id = None
            index += 1
            continue

        if stripped == "---":
            flush_paragraph()
            current_numbering_id = None
            add_horizontal_rule(doc)
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            current_numbering_id = None
            raw_table: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                raw_table.append(lines[index].strip())
                index += 1
            parsed = [split_table_row(row) for row in raw_table if not is_table_separator(row)]
            add_table(doc, parsed, one_page=one_page)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            current_numbering_id = None
            level = len(heading_match.group(1)) - 1
            text = heading_match.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            if not one_page and (
                "Source register" in text
                or text.startswith("17. Final feasibility verdict")
            ):
                p.paragraph_format.page_break_before = True
            add_inline(p, text)
            if level == 1:
                add_bottom_rule(p, color=TEAL, size="6")
            index += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            current_numbering_id = None
            quote_parts = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_parts.append(lines[index].strip().lstrip(">").strip())
                index += 1
            add_callout(doc, " ".join(quote_parts), compact=one_page)
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.+)$", line)
        if bullet_match:
            flush_paragraph()
            current_numbering_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet_match.group(2))
            index += 1
            continue

        number_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if number_match:
            flush_paragraph()
            if current_numbering_id is None:
                current_numbering_id = new_numbering_instance(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, current_numbering_id)
            add_inline(p, number_match.group(2))
            index += 1
            continue

        current_numbering_id = None
        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def set_properties(doc: Document, spec: dict) -> None:
    props = doc.core_properties
    props.title = spec["title"].replace("\n", " ")
    props.subject = spec["subject"]
    props.author = "Lyne"
    props.keywords = "Lyne, Jamaica, budget, feasibility, queue management"
    props.comments = "Evidence cut-off 31 July 2026"


def build_one(spec: dict) -> Path:
    source_text = spec["source"].read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc, one_page=spec["one_page"])
    configure_section(doc.sections[0], one_page=spec["one_page"])
    set_properties(doc, spec)

    if spec["one_page"]:
        add_header_footer(doc.sections[0], spec["short_title"], first_page_blank=True)
        add_one_page_opening(doc, spec)
    else:
        add_header_footer(doc.sections[0], spec["short_title"], first_page_blank=True)
        add_cover(doc, spec)

    render_markdown(doc, source_text, one_page=spec["one_page"])

    for section in doc.sections:
        configure_section(section, one_page=spec["one_page"])
        add_header_footer(
            section,
            spec["short_title"],
            first_page_blank=True,
        )

    spec["output"].parent.mkdir(parents=True, exist_ok=True)
    doc.save(spec["output"])
    return spec["output"]


def main() -> None:
    for spec in DOCS:
        print(build_one(spec))


if __name__ == "__main__":
    main()
